#!/usr/bin/env python
# coding: utf-8

# REFERENCE

# The OLS regression equations and the 2SLS regression equations were created by Angrist, J. D., Battistin, E., \& Vuri, D. (2017) in STATA
# I replicated their regressions in Python

# The regression and instrumental variables are created by Angrist, J. D., Battistin, E., \& Vuri, D. (2017). In a small moment: Class size and moral hazard in the Italian mezzogiorno. American Economic Journal: Applied Economics, 9(4), 216–249. https://doi.org/10.1257/app.20160267

# Data is from the replication package from Angrist, J. D., Battistin, E., \& Vuri, D. (2017b, October 1). Replication data for: In a small moment: Class size and moral hazard in the Italian mezzogiorno. openICPSR. https://www.openicpsr.org/openicpsr/project/113698/version/V1/view

# The regression section of code was assisted by the week One notebooks by Prof. Khazra
# Khazra, N., 2024. Implementation Notebooks, Lecture Material, Week One Notebooks.

# Cross validation was used in this project and the code was assisted by the week Two notebooks by Prof. Khazra
# Khazra, N., 2024. Implementation Notebooks, Lecture Material, Week Two Notebooks.

# The ridge and lasso section of code was assisted by the week Three notebooks by Prof. Khazra
# Khazra, N., 2024. Implementation Notebooks, Lecture Material, Week Three Notebooks.

# The regression trees and ensemble methods section of code was assisted by the week Four notebooks by Prof. Khazra
# Khazra, N., 2024. Implementation Notebooks, Lecture Material, Week Four Notebooks.

# The DAG section of code was assisted by the week Five notebooks by Prof. Khazra
# Khazra, N., 2024. Implementation Notebooks, Lecture Material, Week Five Notebooks.

# The refutation test subsection of the DAG section of code was assisted by the report-week6.pdf by Prof. Khazra
# Khazra, N., 2024. report-week6.pdf, Lecture Material,.

# ChatGPT was used throughout the code file for help with code
# OpenAI. (2022, November 30). Introducing ChatGPT. ChatGPT. https://openai.com/index/chatgpt/

# Each section is marked below where each part of Prof. Khazra, notebooks were used.




# The OLS regression equations and the 2SLS regression equations were created by Angrist, J. D., Battistin, E., \& Vuri, D. (2017) in STATA
# I replicated Angrist, J. D., Battistin, E., \& Vuri, D. (2017) STATA regressions in Python

# The regression section of code was assisted by the week one notebooks by Prof. Khazra
# Khazra, N., 2024. Implementation Notebooks, Lecture Material, Week One Notebooks.


# In[1]:

import numpy as np
import pandas as pd
import matplotlib.pyplot as plt

#get_ipython().system('pip install ISLP -q')
import ISLP

from matplotlib.pyplot import subplots

import statsmodels.api as sm

import statsmodels.formula.api as smf

from statsmodels.stats.outliers_influence import variance_inflation_factor as VIF
from statsmodels.stats.anova import anova_lm

from ISLP import load_data
from ISLP.models import (ModelSpec as MS,
                        summarize,
                        poly)

data=pd.read_stata('smallmo.dta')


# In[ ]:


data.head
data.info
data.describe


# In[ ]:


#look at the columns
data.columns
#see if any NaN
data.isna().sum()


# In[ ]:


#remove the NaN
data.replace([np.inf, not np.inf], np.nan, inplace=True)
data=data.dropna()
#check again to see if no more NaN
data.isna().sum()


# In[ ]:


#summary of data

print(data.describe())

# In[ ]:


#regression where response is math 

#y_math=data['answers_math_std']
#x_math=data.columns.drop('answers_math_std')
#X_math=MS(x_math).fit_transform(data)
#model_math=sm.OLS(y_math, X_math.astype(float))
#results_math=model_math.fit()
#summarize(results_math)


# In[ ]:

# new dataset with out math and italian results
#X_no_score = data.columns.drop(['answers_math_std', 'answers_ital_std', 'schoolid', 'classid','plessoid'])
# Define control variables
#X = MS(X_no_score).fit_transform(data)
# Define dependent variables
#Y = data['answers_math_std']
#model = sm.OLS(Y, X)
#results = model.fit()
#summarize(results)

# In[ ]:


#for trees
from sklearn.tree import (DecisionTreeClassifier as DTC,
                          DecisionTreeRegressor as DTR,
                          plot_tree,
                          export_text)
from sklearn.metrics import (accuracy_score,
                             log_loss)
from sklearn.ensemble import \
     (RandomForestRegressor as RF,
      GradientBoostingRegressor as GBR)
import sklearn.model_selection as skm
#end trees

#start DAG
import numpy as np
from scipy import stats

import matplotlib.pyplot as plt
plt.style.use('fivethirtyeight')

import graphviz
#end DAG

#THIS IS WHERE REPLICATION STARTS
#access rows of data by using the loc[] for data and change to category
data.loc[:, 'survey']=data['survey'].astype('category')
data.loc[:, 'grade']=data['grade'].astype('category')
data.loc[:, 'region']=data['region'].astype('category')
data.loc[:, 'segment']=data['segment'].astype('category')
data.loc[:, 'enrol_ins_snv']=data['enrol_ins_snv'].astype('category')

#create interactions
data['students:segment'] = data['students'] * data['segment']
data['students2:segment'] = data['students2'] * data['segment']

#replicate STATA variables the researchers made
Y = ['answers_math_std', 'answers_ital_std']

X = 'female + m_female + immigrants_broad + m_origin + dad_lowedu + dad_midedu + dad_highedu + mom_unemp + mom_housew + mom_employed + m_dad_edu + m_mom_edu + m_mom_occ'

POLY = 'students + students2 + students:segment + students2:segment + segment'

#I also want to do C(enrol_ins_snv) * C(region) but the code runs forever region has 54 and enrol_ins_snv has 314 this many categories is too much but the results are close to the papers and same direction
FIXED = 'C(survey) + C(grade) + enrol_ins_snv * region'


#MY ADDITION
data['edu_gap'] = abs(data['m_mom_edu'] - data['m_dad_edu'])
data['interaction_effect_eg_clsize'] = data['edu_gap'] * data['clsize_snv']
MYADDITION = 'interaction_effect_eg_clsize + edu_gap'



#string formatting it together like they did in ISLP textbook
CONTROLS = f'{X} + {POLY} + {FIXED} + {MYADDITION}'

#STATA divided class size by 10 after
data['clsize_snv'] = data['clsize_snv']/10



###
#CHAT GPT HELPED FOR THIS CLUSTERED OLS REGRESSION MAINLY LINE cluster_model_math
#PIAZZA recommended to do the smf.ols which helped a lot
model_math = smf.ols(f'answers_math_std ~ clsize_snv + {CONTROLS}', data=data).fit()
#This below line CHATGPT created
cluster_model_math = model_math.get_robustcov_results(cov_type='cluster', groups=data['clu'])
#This above line CHATGPT created
answers_math_std_cluster_model = cluster_model_math.summary()

model_ital = smf.ols(f'answers_ital_std ~ clsize_snv + {CONTROLS}', data=data).fit()
#This below line CHATGPT created
cluster_model_ital = model_ital.get_robustcov_results(cov_type='cluster', groups=data['clu'])
#This above line CHATGPT created
answers_ital_std_cluster_model = cluster_model_ital.summary()
###

print(answers_math_std_cluster_model)
print(answers_ital_std_cluster_model)

#NOW IV PART WITH 2SLS

#math and ital use same IV
#predict clsize_snv with Maimondides Class size rule (clsize_hat)
first_stage = smf.ols(f'clsize_snv ~ clsize_hat + o_math + {CONTROLS}', data=data).fit()
print(first_stage.summary())
#fit values
data['clsize_snv_hat'] = first_stage.fittedvalues

#now math_test_std
#stage two is made using fitted values from clsize_snv_hat
second_stage_math = smf.ols(f'answers_math_std ~ clsize_snv_hat + {CONTROLS}', data=data).fit()
cluster_model_math_2SLS = second_stage_math.get_robustcov_results(cov_type='cluster', groups=data['clu'])
print(cluster_model_math_2SLS.summary())

#now ital_test_std
second_stage_ital= smf.ols(f'answers_ital_std ~ clsize_snv_hat + {CONTROLS}', data=data).fit()
cluster_model_ital_2SLS = second_stage_ital.get_robustcov_results(cov_type='cluster', groups=data['clu'])
print(cluster_model_ital_2SLS.summary())








# RIDGE AND LASSO
# The ridge and lasso section of code was assisted by the week Three notebooks by Prof. Khazra
# Khazra, N., 2024. Implementation Notebooks, Lecture Material, Week Three Notebooks.
#start of ridge regularization

from statsmodels.api import OLS
import sklearn.model_selection as skm
import sklearn.linear_model as skl
from sklearn.preprocessing import StandardScaler
from sklearn.pipeline import Pipeline
#from sklearn.decomposition import PCA, code didnt work with this
from sklearn.cross_decomposition import PLSRegression
#from l0bnb import fit_path, code didnt work with this


#control_vars = 'female + m_female + immigrants_broad + m_origin + dad_lowedu + dad_midedu + dad_highedu + mom_unemp + mom_housew + mom_employed + m_dad_edu + m_mom_edu + m_mom_occ + students + students2 + segment + survey + grade + enrol_ins_snv + region'
#control_vars_list = control_vars.split(' + ')

control_vars = f'{X} + {POLY} + C(grade) + enrol_ins_snv + clsize_hat + o_math + interaction_effect_eg_clsize + edu_gap + C(survey) + region'
control_vars_list = control_vars.split(' + ')
control_vars_list = [var if 'C(' not in var else var.split('C(')[1].split(')')[0] for var in control_vars_list] #from chatgpt need to remove C() to get actual column name


y_math=data['answers_math_std']
x_math=data.columns.drop('answers_math_std')
x_math = data[control_vars_list]
X_math=MS(x_math).fit_transform(data)
model_math=sm.OLS(y_math, X_math.astype(float))
results_math=model_math.fit()
summarize(results_math)

X_math_ridge=X_math.drop('intercept',axis=1)
X_math_ridge = X_math_ridge.astype(float)
X_math_ridge.mean(0)

X_math_ridge=np.asarray(X_math_ridge)
X_math_ridges = X_math_ridge - X_math_ridge.mean(0)[None,:]
X_math_ridge_scale = X_math_ridge.std(0)
X_math_ridges = X_math_ridges / X_math_ridge_scale[None,:]
lambdas = 10**np.linspace(8, -2, 100) / y_math.std()
soln_array = skl.ElasticNet.path(X_math_ridges,
                                 y_math,
                                 l1_ratio=0.,
                                 alphas=lambdas)[1]
soln_array.shape

X_math_ridge = pd.DataFrame(X_math_ridge, columns=X_math.drop('intercept', axis=1).columns)
soln_path = pd.DataFrame(soln_array.T,
                         columns=X_math_ridge.columns,
                         index=-np.log(lambdas))
soln_path.index.name = 'negative log(lambda)'
soln_path

path_fig, ax = subplots(figsize=(8,8))
soln_path.plot(ax=ax, legend=False)
ax.set_xlabel('$-\log(\lambda)$', fontsize=20)
ax.set_ylabel('Standardized coefficients', fontsize=20)
ax.legend(loc='upper left');

import matplotlib.pyplot as plt

plt.savefig("Ridge Regression")
plt.close()

###
### now CV
###
K = 5
kfold = skm.KFold(K,
                  random_state=0,
                  shuffle=True)

#ridge = skl.ElasticNet(alpha=lambdas[59], l1_ratio=0)
from sklearn.linear_model import Ridge
ridge = Ridge()

scaler = StandardScaler(with_mean=True,  with_std=True)
pipe = Pipeline(steps=[('scaler', scaler), ('ridge', ridge)])

param_grid = {'ridge__alpha': lambdas}

grid = skm.GridSearchCV(pipe,
                        param_grid,
                        cv=kfold,
                        scoring='neg_mean_squared_error',
                        n_jobs=-1)
grid.fit(X_math_ridge, y_math)
ridge_optimal_alpha = grid.best_params_['ridge__alpha']
ridge_optimal_be = grid.best_estimator_

mean_test_score = -grid.cv_results_['mean_test_score']
std_test_score = grid.cv_results_['std_test_score'] / np.sqrt(K)

print(grid.cv_results_['mean_test_score'])
print(grid.cv_results_['std_test_score'])

tuned_ridge = pipe.named_steps['ridge'] #new
ridge_fig, ax = subplots(figsize=(12,12))
ax.errorbar(-np.log(lambdas),
            mean_test_score,
            yerr=std_test_score,
            capsize=5,
            markersize=8,
            elinewidth=3)
ax.axvline(-np.log(ridge_optimal_alpha), c='k', ls='--') #new
ax.set_ylim([mean_test_score.min() * 0.9, mean_test_score.max() * 1.1])#mean_test_score is negative not in jupyter range
ax.set_xlabel('$-\log(\lambda)$', fontsize=20)
ax.set_ylabel('Cross-validated MSE', fontsize=20);


plt.savefig("Ridge Regression Cross validation MSE")
plt.close()
### end of ridge




### start of lasso
#control_vars = 'female + m_female + immigrants_broad + m_origin + dad_lowedu + dad_midedu + dad_highedu + mom_unemp + mom_housew + mom_employed + m_dad_edu + m_mom_edu + m_mom_occ + students + students2 + segment + survey + grade + enrol_ins_snv + region'
#control_vars_list = control_vars.split(' + ')

control_vars = f'{X} + {POLY} + C(grade) + enrol_ins_snv + clsize_hat + o_math + interaction_effect_eg_clsize + edu_gap + C(survey) + region'
control_vars_list = control_vars.split(' + ')
control_vars_list = [var if 'C(' not in var else var.split('C(')[1].split(')')[0] for var in control_vars_list] #from chatgpt need to remove C() to get actual column name




y_math=data['answers_math_std']
x_math=data.columns.drop('answers_math_std')
x_math=data[control_vars_list]

X_math=MS(x_math).fit_transform(data)

X_math_lasso = X_math.drop('intercept', axis=1)
X_math_lasso = X_math_lasso.astype(float)

X_math_lasso = np.asarray(X_math_lasso)
X_math_lassos = (X_math_lasso - X_math_lasso.mean(0)[None, :]) / X_math_lasso.std(0)[None, :]

lambdas = 10**np.linspace(8, -2, 100) / y_math.std()

soln_array = skl.ElasticNet.path(X_math_lassos,
                                 y_math,
                                 l1_ratio=1,
                                 alphas=lambdas)[1]
soln_array.shape


X_math_lasso = pd.DataFrame(X_math_lasso, columns=X_math.drop('intercept', axis=1).columns)
soln_path = pd.DataFrame(soln_array.T,
                         columns=X_math_lasso.columns,
                         index=-np.log(lambdas))
soln_path.index.name = 'negative log(lambda)'
soln_path

path_fig, ax = subplots(figsize=(8,8))
soln_path.plot(ax=ax, legend=False)
ax.set_xlabel('$-\log(\lambda)$', fontsize=20)
ax.set_ylabel('Standardized coefficients', fontsize=20)
ax.legend(loc='upper left');

plt.savefig("Lasso Regression")
plt.close()

###CV
from sklearn.linear_model import Lasso
lasso = Lasso()

scaler = StandardScaler(with_mean=True,  with_std=True)
pipe = Pipeline(steps=[('scaler', scaler), ('lasso', lasso)])

param_grid = {'lasso__alpha': lambdas}

grid = skm.GridSearchCV(pipe,
                        param_grid,
                        cv=kfold,
                        scoring='neg_mean_squared_error',
                        n_jobs=-1)
grid.fit(X_math_lasso, y_math)

lasso_optimal_alpha = grid.best_params_['lasso__alpha']
lasso_optimal_be = grid.best_estimator_

mean_test_score = -grid.cv_results_['mean_test_score']
std_test_score = grid.cv_results_['std_test_score'] / np.sqrt(K)

tuned_lasso = pipe.named_steps['lasso'] #new

lasso_fig, ax = subplots(figsize=(12,12))
ax.errorbar(-np.log(lambdas),
            mean_test_score,
            yerr=std_test_score,
            capsize=5,
            markersize=8,
            elinewidth=3)
ax.axvline(-np.log(lasso_optimal_alpha), c='k', ls='--')
ax.set_ylim([mean_test_score.min() * 0.9, mean_test_score.max() * 1.1])#mean_test_score is negative not in jupyter range
ax.set_xlabel('$-\log(\lambda)$', fontsize=20)
ax.set_ylabel('Cross-validated MSE', fontsize=20);

plt.savefig("Lasso Regression Cross validation MSE")
plt.close()

### end of lasso


# END OF RIDGE AND LASSO




# START OF TREES

# The regression trees and ensemble methods section of code was assisted by the week Four notebooks by Prof. Khazra
# Khazra, N., 2024. Implementation Notebooks, Lecture Material, Week Four Notebooks.

control_vars = f'{X} + {POLY} + C(survey) + C(grade) + enrol_ins_snv + region + clsize_hat + o_math + interaction_effect_eg_clsize + edu_gap + clsize_snv'
control_vars_list = control_vars.split(' + ')
control_vars_list = [var if 'C(' not in var else var.split('C(')[1].split(')')[0] for var in control_vars_list] #from chatgpt need to remove C() to get actual column name


y_math=data['answers_math_std']
x_math=data.columns.drop('answers_math_std')
x_math=data[control_vars_list]

X_math=MS(x_math).fit_transform(data)

feature_names = list(X_math.columns)
X_tree = np.asarray(X_math)


(X_train,
 X_test,
 y_train,
 y_test) = skm.train_test_split(X_tree,
                                data['answers_math_std'],
                                test_size=0.3,
                                random_state=0)

reg = DTR(max_depth=3)
reg.fit(X_train, y_train)
ax = subplots(figsize=(12,12))[1]
plot_tree(reg,
          feature_names=feature_names,
          ax=ax);

plt.savefig("Tree")
plt.close()

#cross validation to prune tree

ccp_path = reg.cost_complexity_pruning_path(X_train, y_train)
kfold = skm.KFold(5,
                  shuffle=True,
                  random_state=10)
grid = skm.GridSearchCV(reg,
                        {'ccp_alpha': ccp_path.ccp_alphas},
                        refit=True,
                        cv=kfold,
                        scoring='neg_mean_squared_error')
G = grid.fit(X_train, y_train)

best_ = grid.best_estimator_
print(np.mean((y_test - best_.predict(X_test))**2))

ax = subplots(figsize=(12,12))[1]
plot_tree(G.best_estimator_,
          feature_names=feature_names,
          ax=ax);

plt.savefig("Pruned Tree")
plt.close()

print(np.mean((y_test - reg.predict(X_test))**2))


#bagging

bag_boston = RF(max_features=X_train.shape[1], random_state=0, n_jobs=-1)
bag_boston.fit(X_train, y_train)

ax = subplots(figsize=(8,8))[1]
y_hat_bag = bag_boston.predict(X_test)
ax.scatter(y_hat_bag, y_test)
print(np.mean((y_test - y_hat_bag)**2))

#bag_boston = RF(max_features=X_train.shape[1],
#                n_estimators=500,
#                random_state=0,n_jobs=-1).fit(X_train, y_train)
#y_hat_bag = bag_boston.predict(X_test)
#np.mean((y_test - y_hat_bag)**2)

RF_boston = RF(max_features=6,
               random_state=0, n_jobs=-1).fit(X_train, y_train)
y_hat_RF = RF_boston.predict(X_test)
print(np.mean((y_test - y_hat_RF)**2))

feature_imp = pd.DataFrame(
    {'importance':RF_boston.feature_importances_},
    index=feature_names)

print(feature_imp.sort_values(by='importance', ascending=False))
print(feature_imp.sort_values(by='importance', ascending=False).to_latex())

#boosting

boost_boston = GBR(n_estimators=5000,
                   learning_rate=0.1,
                   max_depth=3,
                   random_state=0,
                   tol=1e-5,
                   n_iter_no_change=20)
boost_boston.fit(X_train, y_train)

test_error = np.zeros_like(boost_boston.train_score_)
for idx, y_ in enumerate(boost_boston.staged_predict(X_test)):
   test_error[idx] = np.mean((y_test - y_)**2)

plot_idx = np.arange(boost_boston.train_score_.shape[0])
ax = subplots(figsize=(8,8))[1]
ax.plot(plot_idx,
        boost_boston.train_score_,
        'b',
        label='Training')
ax.plot(plot_idx,
        test_error,
        'r',
        label='Test')
ax.legend();

plt.savefig("Prediction and Error Plot Boosting")
plt.close()


y_hat_boost = boost_boston.predict(X_test);
print(np.mean((y_test - y_hat_boost)**2))


# END OF TREES





# DAG


# The DAG section of code was assisted by the week Five notebooks by Prof. Khazra
# Khazra, N., 2024. Implementation Notebooks, Lecture Material, Week Five Notebooks.

DAG_control_vars = f'{X} + {POLY} + C(survey) + C(grade) + enrol_ins_snv + region + interaction_effect_eg_clsize + edu_gap'
DAG_control_vars_list = DAG_control_vars.split(' + ')
DAG_control_vars_list = [var if 'C(' not in var else var.split('C(')[1].split(')')[0] for var in DAG_control_vars_list] #from chatgpt need to remove C() to get actual column name


DAG_y_math=data['answers_math_std']
DAG_x_math=data.columns.drop('answers_math_std')

treatment = 'clsize_snv'
instruments = ['clsize_hat', 'o_math']
outcome = 'answers_math_std'

stuff = [treatment] + instruments + [outcome]

DAG_x_math=data[DAG_control_vars_list + stuff]

X_math=MS(DAG_x_math).fit_transform(data)


#NEW DAG
import dowhy
from dowhy import CausalModel

gml_graph = """
graph [
    directed 1
    
    node [
        id "female" 
        label "female"
    ]    
    node [
        id "m_female"
        label "m_female"
    ]
    node [
        id "immigrants_broad"
        label "immigrants_broad"
    ]
    node [
        id "m_origin"
        label "m_origin"
    ]
    node [
        id "dad_lowedu"
        label "dad_lowedu"
    ]
    node [
        id "dad_midedu"
        label "dad_midedu"
    ]
    node [
        id "dad_highedu"
        label "dad_highedu"
    ]
    node [
        id "mom_unemp"
        label "mom_unemp"
    ]
    node [
        id "mom_housew"
        label "mom_housew"
    ]
    node [
        id "mom_employed"
        label "mom_employed"
    ]
    node [
        id "m_dad_edu"
        label "m_dad_edu"
    ]
    node [
        id "m_mom_edu"
        label "m_mom_edu"
    ]
    node [
        id "m_mom_occ"
        label "m_mom_occ"
    ]
    node [
        id "students"
        label "students"
    ]
    node [
        id "students2"
        label "students2"
    ]
    node [
        id "students:segment"
        label "students:segment"
    ]
    node [
        id "students2:segment"
        label "students2:segment"
    ]
    node [
        id "segment"
        label "segment"
    ]
    node [
        id "survey"
        label "survey"
    ]
    node [
        id "grade"
        label "grade"
    ]
    node [
        id "enrol_ins_snv"
        label "enrol_ins_snv"
    ]
    node [
        id "region"
        label "region"
    ]
    node [
        id "clsize_hat"
        label "clsize_hat"
    ]
    node [
        id "clsize_snv"
        label "clsize_snv"
    ]
    node [
        id "answers_math_std"
        label "answers_math_std"
    ]
    node [
        id "o_math"
        label "o_math"
    ]
    node [
        id "interaction_effect_eg_clsize"
        label "interaction_effect_eg_clsize"
    ]
    node [
        id "edu_gap" 
        label "edu_gap"
    ]
    
    
    
    edge [
        source "female"
        target "clsize_snv"
    ]
    edge [
        source "female"
        target "answers_math_std"
    ]
    edge [
        source "m_female"
        target "clsize_snv"
    ]
    edge [
        source "m_female"
        target "answers_math_std"
    ]
    edge [
        source "female"
        target "m_female"
    ]
    edge [
        source "clsize_hat"
        target "clsize_snv"
    ]
    edge [
        source "grade"
        target "clsize_snv"
    ]
    edge [
        source "grade"
        target "answers_math_std"
    ]
    edge [
        source "region"
        target "answers_math_std"
    ]
    edge [
        source "region"
        target "clsize_snv"
    ]
    edge [
        source "enrol_ins_snv"
        target "clsize_snv"
    ]
    edge [
        source "enrol_ins_snv"
        target "answers_math_std"
    ]
    edge [
        source "survey"
        target "answers_math_std"
    ]
    edge [
        source "survey"
        target "clsize_snv"
    ]
    edge [
        source "clsize_snv"
        target "answers_math_std"
    ]
    edge [
        source "m_origin"
        target "immigrants_broad"
    ]
    edge [
        source "dad_lowedu"
        target "m_dad_edu"
    ]
    edge [
        source "dad_midedu"
        target "m_dad_edu"
    ]
    edge [
        source "dad_highedu"
        target "m_dad_edu"
    ]
    edge [
        source "m_dad_edu"
        target "answers_math_std"
    ]
    edge [
        source "segment"
        target "clsize_snv"
    ]
    edge [
        source "segment"
        target "answers_math_std"
    ]
    edge [
        source "students"
        target "students2"
    ]
    edge [
        source "students"
        target "clsize_snv"
    ]
    edge [
        source "students"
        target "answers_math_std"
    ]
    edge [
        source "students2"
        target "answers_math_std"
    ]
    edge [
        source "students2"
        target "clsize_snv"
    ]
    edge [
        source "students:segment"
        target "clsize_snv"
    ]
    edge [
        source "students2:segment"
        target "clsize_snv"
    ]
    edge [
        source "m_mom_occ"
        target "clsize_snv"
    ]
    edge [
        source "m_mom_occ"
        target "answers_math_std"
    ]
    edge [
        source "m_mom_edu"
        target "m_mom_occ"
    ]
    edge [
        source "m_mom_edu"
        target "mom_employed"
    ]
    edge [
        source "m_mom_edu"
        target "answers_math_std"
    ]
    edge [
        source "m_mom_edu"
        target "clsize_snv"
    ]
    edge [
        source "mom_housew"
        target "mom_unemp"
    ]
    edge [
        source "m_mom_edu"
        target "mom_unemp"
    ]
    edge [
        source "mom_unemp"
        target "mom_employed"
    ]
    edge [
        source "students:segment"
        target "answers_math_std"
    ]
    edge [
        source "students2:segment"
        target "answers_math_std"
    ]
    edge [
        source "students"
        target "o_math"
    ]
    edge [
        source "segment"
        target "o_math"
    ]
    edge [
        source "enrol_ins_snv"
        target "o_math"
    ]
    edge [
        source "m_mom_edu"
        target "o_math"
    ]
    edge [
        source "dad_lowedu"
        target "o_math"
    ]
    edge [
        source "dad_midedu"
        target "o_math"
    ]
    edge [
        source "dad_highedu"
        target "o_math"
    ]
    edge [
        source "grade"
        target "o_math"
    ]
    edge [
        source "o_math"
        target "answers_math_std"
    ]
    edge [
        source "m_dad_edu"
        target "o_math"
    ]
    edge [
        source "mom_employed"
        target "o_math"
    ]
    edge [
        source "mom_unemp"
        target "o_math"
    ]
    edge [
        source "students2"
        target "o_math"
    ]
    edge [
        source "region"
        target "o_math"
    ]
    edge [
        source "students2:segment"
        target "o_math"
    ]
    edge [
        source "students:segment"
        target "o_math"
    ]
    edge [
        source "interaction_effect_eg_clsize"
        target "clsize_snv"
    ]
    edge [
        source "interaction_effect_eg_clsize"
        target "answers_math_std"
    ]
    edge [
        source "dad_lowedu"
        target "interaction_effect_eg_clsize"
    ]
    edge [
        source "dad_midedu"
        target "interaction_effect_eg_clsize"
    ]
    edge [
        source "dad_highedu"
        target "interaction_effect_eg_clsize"
    ]
    edge [
        source "m_mom_edu"
        target "interaction_effect_eg_clsize"
    ]
    edge [
        source "m_dad_edu"
        target "interaction_effect_eg_clsize"
    ]
    edge [
        source "students"
        target "interaction_effect_eg_clsize"
    ]
    edge [
        source "enrol_ins_snv"
        target "interaction_effect_eg_clsize"
    ]
    edge [
        source "region"
        target "interaction_effect_eg_clsize"
    ]
    edge [
        source "segment"
        target "interaction_effect_eg_clsize"
    ]
    edge [
        source "students2:segment"
        target "interaction_effect_eg_clsize"
    ]
    edge [
        source "students:segment"
        target "interaction_effect_eg_clsize"
    ]
    edge [
        source "dad_lowedu"
        target "edu_gap"
    ]
    edge [
        source "dad_midedu"
        target "edu_gap"
    ]
    edge [
        source "dad_highedu"
        target "edu_gap"
    ]
    edge [
        source "m_dad_edu"
        target "edu_gap"
    ]
    edge [
        source "m_mom_edu"
        target "edu_gap"
    ]
    edge [
        source "edu_gap"
        target "interaction_effect_eg_clsize"
    ]
    edge [
        source "edu_gap"
        target "answers_math_std"
    ]
]
"""
# This subsection of the DAG section of code was assisted by the report-week6.pdf by Prof. Khazra
# Khazra, N., 2024. report-week6.pdf, Lecture Material,.

backdoors = ['m_mom_edu','m_mom_occ','female', 'm_female', 'grade', 'region', 'enrol_ins_snv', 'survey', 'segment','students','students2', 'students2:segment', 'students:segment', 'interaction_effect_eg_clsize']

model=CausalModel(
    data=DAG_x_math,
    treatment='clsize_snv',
    outcome='answers_math_std',
    instruments=['clsize_hat', 'o_math'],
    common_causes=backdoors,
    graph=gml_graph
)
model.view_model()

print(DAG_x_math.columns)




print("SPLIT IDENTIFY ESTIMAND")

#identify the estimand
estimand = model.identify_effect()
print(estimand)




print("SPLIT ESTIMATE CAUSAL EFFECT")

#estiamte the causal effect / obtain estimates
estimate = model.estimate_effect(
    identified_estimand=estimand,
    method_name='iv.instrumental_variable')

print(f'Estimate of causal effect (linear regression): {estimate.value}')



print("SPLIT REFUTATION TESTS")

#run refutation tests
refute_subset = model.refute_estimate(
    estimand=estimand,
    estimate=estimate,
    method_name="data_subset_refuter",
    subset_fraction=0.4)

print(refute_subset)

print("SPLIT")

#OLD DAG
"""""
# FULL DAG
treatment = 'clsize_snv'
instruments = 'clsize_hat'
outcome = 'answers_math_std'




from dowhy import CausalModel
model=CausalModel(
    data=data,
    treatment=treatment,
    outcome=outcome,
    common_causes=DAG_control_vars_list,
    instruments=instruments,
    effect_modifiers=['female', 'm_female', 'm_mom_edu', 'm_dad_edu']
)
model.view_model()



# DAG WITH SOME CONFOUDNERS REMOVED
DAG_Small_control_vars_list = DAG_control_vars_list.copy()

DAG_Small_control_vars_list.remove('region')
DAG_Small_control_vars_list.remove('immigrants_broad')
DAG_Small_control_vars_list.remove('mom_housew')
DAG_Small_control_vars_list.remove('mom_unemp')
DAG_Small_control_vars_list.remove('female')


from dowhy import CausalModel
model=CausalModel(
    data=data,
    treatment=treatment,
    outcome=outcome,
    common_causes=DAG_Small_control_vars_list,
    instruments=instruments,
    effect_modifiers=None
)
model.view_model()


#LOG POSE DAG

DAG_LOG_POSE_control_vars_list = DAG_control_vars_list.copy()
DAG_LOG_POSE_control_vars_list.append('north')
DAG_LOG_POSE_control_vars_list.append('centre')
DAG_LOG_POSE_control_vars_list.append('south')


from dowhy import CausalModel
model=CausalModel(
    data=data,
    treatment=treatment,
    outcome=outcome,
    common_causes=DAG_LOG_POSE_control_vars_list,
    instruments=instruments,
    effect_modifiers=['female', 'm_female', 'north', 'south', 'centre']
)
model.view_model()
"""""
#END OF OLD DAG




#put the summary statistics table and or a regression table into word doc for important variables
from docx import Document
from docx.shared import Inches
document = Document()


#This is written report 1

#1. Dataset Acquisition

document.add_heading('1. Dataset Acquisition', level=1)
document.add_paragraph('The dataset comes from the American Economic Journal: Applied Economics by the '
                       ' American Economic Association. The paper is titled'
                       ' "In a Small Moment: Class Size and Moral Hazard in the Italian Mezzogiorno".'
                       ' The authors are Joshua D. Angrist, Erich Battistin'
                       ' and Daniela Vuri. The dataset comes from the INVALSI standardized test results of second and'
                       ' fifth graders in Italy. The type of data in this dataset is observational data obtained'
                       ' from administrative sources. The data is from 2009-2012 and includes key variables'
                       ' such as test scores for math and Italian, class size, gender, parental employment'
                       ' and parental education. I have successfully downloaded the dataset and the filename is '
                       ' "smallmo.dta" and the link is '
                       ' "https://www.openicpsr.org/openicpsr/project/113698/version/V1/view?path=/openicpsr/113698/fcr:versions/V1/smallmo.dta&type=file".')

#2. Understanding the Paper’s Identification Strategy and Assumptions
document.add_heading('2. Understanding the Paper’s Identification Strategy and Assumptions', level=1)
document.add_paragraph('The identification strategy used in the journal paper by the authors was '
                       ' instrumental variables. The authors used the Maimonides Rule as an instrumental variable'
                       ' to identify the causal impact of class size on student grades. The Maimonides Rule '
                       ' affects the class size as it is used to determine the class size since if '
                       ' the class size exceeds a certain cutoff a new class is made. There is a problem'
                       ' since class size affects the math and Italian grades of a student so there is an'
                       ' endogeneity problem, but the Maimonides Rule helps solve this problem. There are three key '
                       ' assumptions made by the authors. The first is'
                       ' that class size reductions are exogenously determined by the Maimonides Rule. The second is'
                       ' that teachers manipulate test marks more in larger classes and depends on the teacher'
                       ' behaviour. The third is class size cutoff rules are followed. Since I will be replicating'
                       ' the causal analysis done in this paper then apply machine learning techniques to it it is '
                       ' important for me to understand their strategy and its assumptions. Using the'
                       ' Maimonides Rule I can replicate the causal effect of class size on a students grade. After'
                       ' I will be able to apply machine learning techniques after to provide better predictions into'
                       ' the causal relationships. Using machine learning is very useful when it comes to missing'
                       ' data which is often a big problem in observational studies and may be a problem in this study.')

#3. Summary Statistics of Important Variables
document.add_heading('3. Summary Statistics of Important Variables', level=1)
document.add_paragraph('The key variables in the dataset are female, m_origin, dad_midedu, mom_employed'
                       ' enrol_sch_snv, immigrants_broad, answers_math_pct, answers_ital_pct, clsize_snv,'
                       ' our_CHEAT_math, our_CHEAT_ital, answers_math_std, answers_ital_std. The m infront of the'
                       ' variables just mean that they are the version of the variable with no missing data. '
                       ' Below in Table 2 the means, medians, standard deviations, minimums and maximums'
                       ' are displayed for the key variables. Almost 49% of the students are female, '
                       ' 25.5% of dads have graduated high school and '
                       ' 45% of moms are employed. Table 1 shows the same descriptive statistics for the '
                       ' variables in the identification strategy. The class size variable is the central variable'
                       ' for the causal analysis.'
                       ' The class size has a mean of around 20 students,'
                       ' the smallest class is 11 students and the largest class is 30 students. The dataset'
                       ' is of high quality as it has well defined labels and some missing values. '
                       ' No anomalies have been found so far such as no class was below 10 or above 30 so the'
                       ' Maimonides Rule is holding. As well the class sizes are evenly spread. Table 3 and Table 4 '
                       ' show regression outputs for the response variables math scores and Italian scores.')


#4. Research Question and Planned Analysis
document.add_heading('4. Research Question and Planned Analysis', level=1)
document.add_paragraph('Since I will be replicating the paper I will be using a similar research question.'
                       ' I will replicate the same causal analysis from the paper but I extend the analysis by applying'
                       ' machine learning techniques to the question. I will use the Maimonides Rule as '
                       ' my instrumental variable for class size. The research question I will be exploring will be'
                       ' "The Impact of Class Size on Student Grades with Machine Learning Techniques". '
                       ' This is an important question to answer since it can inform education '
                       ' policy decisions on class size. This paper will contribute to'
                       ' existing literature since using machine learning techniques may enhance prediction'
                       ' accuracy and may help with the missing data. This also contributes to existing literature'
                       ' since causal machine learning is an evolving field with less research compared to '
                       ' econometric methods. The exact machine learning technique I will be using has not been'
                       ' finalized as of yet since the class will go over machine learning techniques in week 3.')

###
###
###
#summary statistics table for all variables
#document.add_heading('Summary Statistics Table', level=1)
#document.add_paragraph(data.describe().round(3).to_string())
###
###
###
document.add_heading('Table 1: Summary Statistics Table for Class Size', level=1)
#document.add_paragraph((data['clsize_snv']*10).describe().round(3).to_string())
data['clsize_snv_10'] = data['clsize_snv']*10
latex_table_class_size = data[['clsize_snv_10', 'edu_gap', 'interaction_effect_eg_clsize']].describe().round(4).T.to_latex()
document.add_paragraph(latex_table_class_size)

#make columns for key variables
key_variables = ['female', 'm_origin', 'dad_midedu', 'mom_employed', 'enrol_sch_snv', 'immigrants_broad', 'answers_math_pct', 'answers_ital_pct', 'clsize_snv', 'our_CHEAT_math', 'our_CHEAT_ital', 'answers_math_std', 'answers_ital_std', 'clsize_hat', 'o_math', 'interaction_effect_eg_clsize','edu_gap']
document.add_heading('Table 2: Summary Statistics Table Key Variables', level=1)
#document.add_paragraph(data[key_variables].describe().round(3).to_string())
#latex_table_key_variables = data[key_variables].describe().round(4).to_latex()
sum_stats = data[key_variables].describe().round(4)
transposed_sum_stats = sum_stats.T
latex_table_key_variables = transposed_sum_stats.to_latex()

document.add_paragraph(latex_table_key_variables)

document.add_heading('Table 3: Regression Table for Math Scores', level=1)
#document.add_paragraph(answers_math_std_cluster_model.as_text())
latex_math_reg = answers_math_std_cluster_model.as_latex()
document.add_paragraph(latex_math_reg)

document.add_heading('Table 4: Regression Table for Italian Scores', level=1)
#document.add_paragraph(answers_ital_std_cluster_model.as_text())
latex_ital_reg = answers_ital_std_cluster_model.as_latex()
document.add_paragraph(latex_ital_reg)

document.save('ECO2425 Project.docx')



#report week 3


X3 = ['clsize_snv','female', 'm_female', 'immigrants_broad', 'm_origin',
                         'dad_lowedu', 'dad_midedu', 'dad_highedu',
                         'mom_unemp', 'mom_housew', 'mom_employed',
                         'm_dad_edu', 'm_mom_edu', 'm_mom_occ', 'enrol_ins_snv', 'segment','students','students2', 'students2:segment', 'students:segment', 'interaction_effect_eg_clsize', 'edu_gap']

week3_reg_table = cluster_model_math.summary2().tables[1]

week3_reg_table.loc['clsize_snv', 'Coef.'] *= 10
week3_reg_table.loc['clsize_snv', 'Std.Err.'] *= 10


###
###
###
available_columns = week3_reg_table.index
categorical_vars = [col for col in available_columns if col.startswith('C(grade)') or
                    col.startswith('C(region)') or col.startswith('C(survey)')]
X3_new = X3 + categorical_vars
X3_f = [col for col in X3_new if col in available_columns]
###
###
###

week3_reg_table_imp = week3_reg_table.loc[X3_f]


latex_output = week3_reg_table_imp.to_latex()

from docx.enum.text import WD_ALIGN_PARAGRAPH

title = document.add_paragraph('The Effect of Class Size on Grades')
name = document.add_paragraph('Eliott Monkman')

title.alignment = WD_ALIGN_PARAGRAPH.CENTER
name.alignment = WD_ALIGN_PARAGRAPH.CENTER

document.add_heading('1 OLS', level=1)
document.add_paragraph(
    "answers_math_std = β₀ + β₁ * clsize_snv + β₂ * female + β₃ * m_female + "
    "β₄ * immigrants_broad + β₅ * m_origin + β₆ * dad_lowedu + β₇ * dad_midedu + "
    "β₈ * dad_highedu + β₉ * mom_unemp + β₁₀ * mom_housew + β₁₁ * mom_employed + "
    "β₁₂ * m_dad_edu + β₁₃ * m_mom_edu + β₁₄ * m_mom_occ + ε"
)
document.add_paragraph(latex_output)
document.add_paragraph('Only key variables are displayed other predictors are omitted for readability')

document.add_heading('2.1 Ridge', level=1)
document.add_picture("C:/Users/pengu/Downloads/Ridge Regression.png", width=Inches(4), height=Inches(3))

document.add_heading('2.2 CV MSE', level=1)
document.add_picture("C:/Users/pengu/Downloads/Ridge Regression Cross validation MSE.png", width=Inches(4), height=Inches(3))
document.add_paragraph(f' lambda = {ridge_optimal_alpha}')
document.add_paragraph(f'-log(lambda) = {-np.log(ridge_optimal_alpha)}')

document.add_heading('3.1 Lasso', level=1)
document.add_picture("C:/Users/pengu/Downloads/Lasso Regression.png", width=Inches(3), height=Inches(3))

document.add_heading('3.2 CV MSE', level=1)
document.add_picture("C:/Users/pengu/Downloads/Lasso Regression Cross validation MSE.png", width=Inches(3), height=Inches(3))
document.add_paragraph(f' lambda = {lasso_optimal_alpha}, -log(lambda) = {-np.log(lasso_optimal_alpha)}')

document.add_paragraph('In the OLS regression some coefficients are really high like the 3 levels of fathers education,'
                       ' this may be due to endogeneity which the instrumental variable will fix. '
                       ' Lasso has a smaller lambda so lasso may'
                       ' be more effective in reducing impact of irrelevant predictors so OLS may overfit, this helps'
                       ' support my decision of not including the full factorial like the paper did since it only caused'
                       ' the replication to be off by around 0.002 and ridge supports that there is high collinearity.')

document.save('ECO2425 Project Week 3 Report.docx')




document = Document()

document.add_heading('First Stage F_Value', level=1)
first_stage_fstat = first_stage.fvalue
document.add_paragraph(str(first_stage_fstat))

document.add_heading('Hausman Test', level=1)
#hausman_test_check = np.dot(np.dot((cluster_model_math.params-cluster_model_math_2SLS.params).T, np.linalg.pinv((cluster_model_math.cov_params()-cluster_model_math_2SLS.cov_params()))), (cluster_model_math.params-cluster_model_math_2SLS.params))
hausman_test_check = np.dot(np.dot((cluster_model_math_2SLS.params-cluster_model_math.params).T, np.linalg.pinv((cluster_model_math_2SLS.cov_params()-cluster_model_math.cov_params()))), (cluster_model_math_2SLS.params-cluster_model_math.params))


hausman_test_p_value = 1 - stats.chi2.cdf(hausman_test_check, len(cluster_model_math.params))
document.add_paragraph(str(hausman_test_check))
document.add_paragraph(str(hausman_test_p_value))


first_state_latex = first_stage.summary().as_latex()
cluster_model_math_2SLS_latex = cluster_model_math_2SLS.summary().as_latex()
cluster_model_ital_2SLS_latex = cluster_model_ital_2SLS.summary().as_latex()

document.add_heading('First Stage', level=1)
document.add_paragraph(first_state_latex)

document.add_heading('2SLS Math', level=1)
document.add_paragraph(cluster_model_math_2SLS_latex)

document.add_heading('2SLS Italian', level=1)
document.add_paragraph(cluster_model_ital_2SLS_latex)


document.save('ECO2425 Project IV OUTPUT.docx')